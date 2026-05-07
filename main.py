import streamlit as st
import g4f
from PIL import Image
import pytesseract
from docx import Document
from io import BytesIO
import fitz
import os

# --- 1. CẤU HÌNH HỆ THỐNG & GIAO DIỆN ---
st.set_page_config(
    page_title="MATH PRO AI - CHÁNH PHÚC", 
    page_icon="📐", 
    layout="wide"
)

# Cấu hình Tesseract cho chạy máy cá nhân (Local)
if os.path.exists(r'C:\Program Files\Tesseract-OCR\tesseract.exe'):
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# --- 2. CÁC HÀM XỬ LÝ CHÍNH ---

def ask_ai(prompt):
    """Gửi yêu cầu tới AI để xử lý nội dung"""
    try:
        return g4f.ChatCompletion.create(
            model="gpt-4", 
            messages=[{"role": "user", "content": prompt}]
        )
    except:
        return "AI hiện đang quá tải, thầy/cô vui lòng thử lại sau giây lát."

def export_to_word(text_content):
    """Xuất văn bản sang file Word hỗ trợ công thức Toán học"""
    doc = Document()
    doc.add_heading('TÀI LIỆU TOÁN HỌC - MATH PRO AI', 0)
    
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. GIAO DIỆN MENU BÊN TRÁI ---
with st.sidebar:
    st.title("🎓 Math Pro AI 2.0")
    st.markdown("---")
    choice = st.radio("Chức năng hệ thống:", [
        "🏠 Trang chủ", 
        "📸 Quét Ảnh/PDF sang Word", 
        "📝 Soạn Đề & Tạo Bài Tập", 
        "✨ Chuyển Văn Bản AI sang Word",
        "⚙️ Chuẩn Hóa LaTeX"
    ])
    st.markdown("---")
    st.info("Phát triển bởi: **CHÁNH PHÚC**")

# --- 4. CHI TIẾT CÁC CHỨC NĂNG ---

if choice == "🏠 Trang chủ":
    st.header("🚀 Trợ Lý Công Nghệ Toán Học Thông Minh")
    st.markdown("""
    Hệ thống hỗ trợ giáo viên tối ưu hóa công việc:
    - **Quét đề:** Chuyển ảnh chụp đề bài thành văn bản Word.
    - **Tạo bài tập:** Tải file mẫu lên để AI tạo các câu hỏi tương tự.
    - **Định dạng:** Chuyển đổi công thức từ AI sang Word chuẩn.
    """)
    st.image("https://img.freepik.com/free-vector/mathematics-concept-illustration_114360-3972.jpg", width=500)

elif choice == "📸 Quét Ảnh/PDF sang Word":
    st.subheader("🖼️ Trích xuất dữ liệu từ hình ảnh")
    up_file = st.file_uploader("Tải lên ảnh đề bài hoặc file PDF", type=["png", "jpg", "jpeg", "pdf"])
    
    if up_file:
        with st.spinner("Đang nhận diện ký tự..."):
            if up_file.type == "application/pdf":
                with fitz.open(stream=up_file.read(), filetype="pdf") as doc:
                    raw_text = "".join([page.get_text() for page in doc])
            else:
                img = Image.open(up_file)
                st.image(img, caption="Ảnh nguồn", width=400)
                raw_text = pytesseract.image_to_string(img, lang='vie')

        st.success("Đã quét xong!")
        edited_text = st.text_area("Văn bản quét được:", value=raw_text, height=200)
        
        if st.button("🪄 AI Làm sạch & Xuất Word"):
            with st.spinner("AI đang định dạng công thức..."):
                final_res = ask_ai(f"Chỉnh sửa lỗi chính tả và chuyển các công thức toán về dạng LaTeX nằm trong dấu $: {edited_text}")
                st.markdown(final_res)
                st.download_button("📥 Tải file Word chuẩn", export_to_word(final_res), "KetQua_Quet.docx")

elif choice == "📝 Soạn Đề & Tạo Bài Tập":
    st.subheader("✍️ Soạn câu hỏi & Tạo bài tập tương tự")
    
    source_type = st.radio("Chọn nguồn dữ liệu:", ["Nhập chủ đề", "Tải file bài tập mẫu (.docx, .pdf)"], horizontal=True)
    context_text = ""
    
    if source_type == "Nhập chủ đề":
        col1, col2 = st.columns(2)
        with col1:
            topic = st.text_input("Chủ đề bài học:", placeholder="VD: Logarit lớp 12")
        with col2:
            num_q = st.number_input("Số lượng câu:", 1, 20, 5)
        context_text = f"về chủ đề {topic}"
    else:
        uploaded_sample = st.file_uploader("Tải file bài tập mẫu để AI bắt chước:", type=["docx", "pdf"])
        num_q = st.number_input("Số lượng câu muốn tạo thêm:", 1, 20, 5)
        
        if uploaded_sample:
            with st.spinner("Đang đọc file mẫu..."):
                if uploaded_sample.type == "application/pdf":
                    with fitz.open(stream=uploaded_sample.read(), filetype="pdf") as doc:
                        context_text = "dựa trên nội dung các bài tập sau: " + "".join([page.get_text() for page in doc])
                else:
                    from docx import Document as ReadDoc
                    doc_read = ReadDoc(uploaded_sample)
                    context_text = "dựa trên nội dung các bài tập sau: " + "\n".join([p.text for p in doc_read.paragraphs])
            st.success("Đã đọc file mẫu thành công!")

    q_style = st.selectbox("Dạng đề mong muốn:", ["Trắc nghiệm 4 lựa chọn", "Trắc nghiệm Đúng/Sai", "Tự luận thay số"])

    if st.button("🔥 Bắt đầu tạo bài tập"):
        if context_text == "":
            st.warning("Vui lòng nhập chủ đề hoặc tải file mẫu.")
        else:
            with st.spinner("AI đang phân tích và soạn đề..."):
                prompt = f"Hãy tạo {num_q} câu hỏi {q_style} {context_text}. Yêu cầu: Giữ đúng mức độ kiến thức, có đáp án và giải thích, các công thức toán dùng dấu $...$."
                res = ask_ai(prompt)
                st.markdown(res)
                st.download_button("📥 Tải bộ đề mới (.docx)", export_to_word(res), "BaiTap_TuongTu.docx")

elif choice == "✨ Chuyển Văn Bản AI sang Word":
    st.subheader("📝 Chuyển công thức AI sang Word chuẩn")
    ai_input = st.text_area("Nội dung từ AI:", height=300, placeholder="Dán nội dung copy từ ChatGPT/Gemini tại đây...")
    
    if st.button("📦 Đóng gói file Word"):
        if ai_input:
            with st.spinner("Đang định dạng..."):
                refined = ask_ai(f"Đảm bảo tất cả các công thức toán trong đoạn sau nằm trong cặp dấu $...$: {ai_input}")
                st.markdown(refined)
                st.download_button("📥 Tải file Word", export_to_word(refined), "Document_Math.docx")

elif choice == "⚙️ Chuẩn Hóa LaTeX":
    st.subheader("🪄 Trình chuẩn hóa mã LaTeX")
    text_math = st.text_area("Nhập văn bản chứa toán học:")
    if st.button("Chuyển sang LaTeX"):
        res = ask_ai(f"Chuyển đoạn văn bản sau sang mã LaTeX chuẩn: {text_math}")
        st.code(res, language="latex")
        st.markdown(res)